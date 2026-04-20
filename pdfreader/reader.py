from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document 
from pypdf import PdfReader


QUESTION_START_RE = re.compile(r"(?m)^(?:Q\s*)?(\d{1,2})(?:[.)]|(?=\s+[A-Za-z]))\s*")
OPTION_RE = re.compile(r"(?i)(?<![A-Za-z])\(?([A-E])\)?[.)]\s*")
QUESTION_LINE_RE = re.compile(r"^(?:Q\s*)?(\d{1,2})(?:[.)]|(?=\s+[A-Za-z]))\s*(.*)$")
INSTRUCTION_RANGE_RE = re.compile(r"Qs?\.?\s*\)?\s*(\d+)\s*(?:to|-)\s*(\d+)", re.IGNORECASE)
ANSWER_RE = re.compile(r"(?i)^(?:Answer|Ans):\s*([A-E](?:\s*,\s*[A-E])*)(?:\s|$)", re.MULTILINE)


def extract_pdf_text(pdf_path: Path) -> str:
    reader = PdfReader(str(pdf_path))
    pages: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"--- Page {index} ---\n{text.strip()}")

    return "\n\n".join(pages).strip()


def extract_docx_text(docx_path: Path) -> str:
    document = Document(str(docx_path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" ".join(cells))

    return "\n".join(parts).strip()


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf_text(file_path)
    if suffix == ".docx":
        return extract_docx_text(file_path)

    raise ValueError(f"Unsupported file type: {file_path.suffix}. Use .pdf or .docx")


def normalize_text(text: str) -> str:
    cleaned = text.replace("\u2026", "...")
    cleaned = cleaned.replace("â€¦", "...")

    lines: list[str] = []
    for raw_line in cleaned.splitlines():
        line = raw_line.strip()
        if not line:
            lines.append("")
            continue
        if re.fullmatch(r"--- Page \d+ ---", line):
            continue
        if re.fullmatch(r"[a-z-]{1,20}", line):
            continue
        lines.append(line)

    cleaned = "\n".join(lines)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def parse_options(options_text: str) -> dict[str, str]:
    matches = list(OPTION_RE.finditer(options_text))
    options: dict[str, str] = {}

    for index, match in enumerate(matches):
        label = match.group(1).upper()
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(options_text)
        value = re.sub(r"\s+", " ", options_text[start:end]).strip(" :")
        options[label] = value

    return options


def parse_answers(answer_text: str) -> list[str] | None:
    """Parse answer string and return list of answers (single or multiple)"""
    if not answer_text:
        return None
    
    # Split by comma and clean up
    answers = [ans.strip().upper() for ans in answer_text.split(",")]
    # Filter out empty strings and validate answers are single letters
    answers = [ans for ans in answers if ans and len(ans) == 1 and ans.isalpha()]
    
    return answers if answers else None


def finalize_question(
    questions: list[dict[str, object]],
    current: dict[str, object] | None,
) -> None:
    if not current or len(current["options"]) < 2:
        return

    question: dict[str, object] = {
        "question_number": current["question_number"],
        "question": re.sub(r"\s+", " ", str(current["question"])).strip(),
        "options": current["options"],
    }
    if current.get("instruction"):
        question["instruction"] = current["instruction"]
    if current.get("answers"):
        question["answers"] = current["answers"]
    questions.append(question)


def parse_instruction_range(instruction: str) -> tuple[int, int] | None:
    match = INSTRUCTION_RANGE_RE.search(instruction)
    if not match:
        return None

    start = int(match.group(1))
    end = int(match.group(2))
    if start > end:
        start, end = end, start
    return start, end


def parse_questions(text: str) -> list[dict[str, object]]:
    normalized = normalize_text(text)
    lines = [line.strip() for line in normalized.splitlines() if line.strip()]
    questions: list[dict[str, object]] = []
    current: dict[str, object] | None = None
    active_instruction = ""
    active_instruction_range: tuple[int, int] | None = None
    pending_instruction: list[str] = []

    for line in lines:
        # Check if this line is an answer line (Answer: or Ans:)
        answer_line_check = re.match(r"(?i)^(?:Answer|Ans):\s*(.*)$", line)
        if answer_line_check and current is not None:
            answer_text = answer_line_check.group(1).strip()
            
            if not answer_text:
                # Empty answer line - log but don't crash
                print(f"WARNING: Question {current.get('question_number')} has empty Answer line")
            else:
                # Try to parse answers
                answers = parse_answers(answer_text)
                if answers:
                    current["answers"] = answers
                else:
                    # Answer line exists but no valid answers found
                    print(f"WARNING: Question {current.get('question_number')} - Answer line has invalid format: '{answer_text}'")
            continue

        if "DIRECTIONS" in line and not line.upper().startswith("DIRECTIONS"):
            line, trailing = line.split("DIRECTIONS", 1)
            line = line.strip()
            if trailing.strip():
                finalize_question(questions, current)
                current = None
                pending_instruction = [f"DIRECTIONS{trailing}".strip()]

        if line.upper().startswith("DIRECTIONS"):
            finalize_question(questions, current)
            current = None
            pending_instruction = [line]
            continue

        if pending_instruction and not QUESTION_LINE_RE.match(line):
            pending_instruction.append(line)
            continue

        question_match = QUESTION_LINE_RE.match(line)
        if question_match:
            finalize_question(questions, current)
            if pending_instruction:
                active_instruction = " ".join(pending_instruction).strip()
                active_instruction_range = parse_instruction_range(active_instruction)
                pending_instruction = []
            question_number = int(question_match.group(1))
            if active_instruction_range and not (
                active_instruction_range[0] <= question_number <= active_instruction_range[1]
            ):
                active_instruction = ""
                active_instruction_range = None
            current = {
                "question_number": question_number,
                "question": question_match.group(2).strip().lstrip("). "),
                "options": {},
                "instruction": active_instruction,
            }
            continue

        if pending_instruction:
            active_instruction = " ".join(pending_instruction).strip()
            active_instruction_range = parse_instruction_range(active_instruction)
            pending_instruction = []

        if current is None:
            continue

        if OPTION_RE.search(line):
            current["options"].update(parse_options(line))
        else:
            current["question"] = f"{current['question']} {line}".strip()

    finalize_question(questions, current)

    return questions


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract text from a PDF or DOCX file.")
    parser.add_argument(
        "input_file",
        nargs="?",
        default="aptitude.pdf",
        help="Path to the input file. Supports .pdf and .docx. Defaults to aptitude.pdf",
    )
    parser.add_argument(
        "-o",
        "--output",
        help="Optional output text file path. If omitted, text is printed.",
    )
    parser.add_argument(
        "--json-output",
        help="Optional JSON output path for parsed questions and options.",
    )
    args = parser.parse_args()

    input_path = Path(args.input_file)
    if not input_path.is_absolute():
        input_path = Path(__file__).resolve().parent / input_path

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    text = extract_text(input_path)
    questions = parse_questions(text)

    if args.json_output:
        json_path = Path(args.json_output)
        if not json_path.is_absolute():
            json_path = Path(__file__).resolve().parent / json_path
        json_path.write_text(json.dumps(questions, indent=2, ensure_ascii=False), encoding="utf-8")
        print(f"Parsed JSON saved to: {json_path}")
        return

    if args.output:
        output_path = Path(args.output)
        if not output_path.is_absolute():
            output_path = Path(__file__).resolve().parent / output_path
        output_path.write_text(text, encoding="utf-8")
        print(f"Extracted text saved to: {output_path}")
        return

    print(text)


if __name__ == "__main__":
    main()
